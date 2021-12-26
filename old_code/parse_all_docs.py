import pandas as pd
from docx import Document
import os


def clean_linebreaks(text):
    """
    returns *text* cleaned from sequences of linebreaks
    """
    text = [x for x in text.split('\n') if x != '']
    return '\n'.join(text)


def parse_tables(document):
    res = []
    for table in document.tables:
        text_in_table = ['Таблица: ']
        for column in table.columns:
            text_in_table.append('; '.join([x.text for x in column.cells]))
        res.append(' / '.join(text_in_table))
    return '\n'.join(res)
    # res = []
    # for table in document.tables:
        # res_table = ''
        # res_table += ('<table>')

        # for column in table.columns:
            # res_table += ('<tr><td>')
            # res_table += ('</td><td>'.join([x.text for x in column.cells]))
            # res_table += ('</td></tr>')
        # res_table += ('</table>')
        # res.append(res_table)
    # return '\n'.join(res)


def parse_all(corpus_path):
    document_texts = {}

    for doc_name in os.listdir(corpus_path):
        print(doc_name)
        doc_path = corpus_path + doc_name
        doc = Document(doc_path)

        a = []
        for para in doc.paragraphs:
            a.append(para.text)
            # a.append('\n')
        a.append(parse_tables(doc))

        doc_text = '\n'.join(a)
        document_texts[doc_name] = doc_text    

    df = pd.DataFrame({
        'doc_name': document_texts.keys(),
        'doc_text': document_texts.values()
        })

    df.doc_text = df.doc_text.apply(clean_linebreaks)
    df.to_excel('output/text_in_docs.xlsx', index=False)


if __name__=="__main__":
    parse_all('data/instructions/')
